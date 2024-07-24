import os
import qrcode
from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import time

def generate_qr_code(data, filename):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=1,
    )
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(filename)

def create_labels_with_qr_codes(data_list, rows, cols, output_filename, label_width, label_height, cube_image_path, single_run = None):
    document = Document()

    sections = document.sections
    for section in sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0)

    table = document.add_table(rows=rows, cols=cols)
    
    # Ajuster les dimensions des cellules pour qu'elles correspondent à la taille des étiquettes
    table.autofit = False
    for row in table.rows:
        row.height = Cm(label_height)
        for cell in row.cells:
            cell.width = Cm(label_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Insérer les QR codes, les images de cubes et le texte dans les cellules du tableau
    
    data_index = 0

    if single_run == True:
        i, j = 0, 0
        while i < 1 or i > 13:
            try:
                i = int(input('Row number (from 1 to 13) : '))
            except ValueError:
                pass
        while j < 1 or j > 5:
            try:
                    j = int(input('Column number (from 1 to 5) : '))
            except ValueError:
                pass

        filename = f'qrcodes/qrcode_{data_index}.png'
        generate_qr_code(data_list[data_index], filename)
        cell = table.cell(i-1, j-1)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.font.size = Pt(4)

        # Ajouter le QR code
        run.add_picture(filename, width=Cm(label_width * 0.4))
        
        # Ajouter l'image du logo Pichon
        run.add_picture(cube_image_path, width=Cm(label_width * 0.2))
        
        # Ajouter le texte
        run.add_text(str(data_list[data_index]))
        print(data_list[data_index])
    else:
        for i in range(rows):
            for j in range(cols):
                if data_index < len(data_list):
                    filename = f'qrcodes/qrcode_{data_index}.png'
                    generate_qr_code(data_list[data_index], filename)
                    cell = table.cell(i, j)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.font.size = Pt(4)

                    # Ajouter le QR code
                    run.add_picture(filename, width=Cm(label_width * 0.4))
                    
                    # Ajouter l'image du logo Pichon
                    run.add_picture(cube_image_path, width=Cm(label_width * 0.2))
                    
                    # Ajouter le texte
                    run.add_text(str(data_list[data_index]))
                    
                    data_index += 1

    document.save(output_filename)

def gen_qrcodes_bulk(table, starting_index, single_run = None):
    if table == 'Computers':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'Computers,{i}' for i in range(starting_index,starting_index + 65) ]
        
    elif table == 'Phones':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'Phones,{i}' for i in range(starting_index,starting_index + 65) ]
    elif table == 'Screens':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'Screens,{i}' for i in range(starting_index,starting_index + 65) ]
    elif table == 'Printers':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'Printers,{i}' for i in range(starting_index,starting_index + 65) ]
    elif table == 'ExternalDrives':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'ExternalDrives,{i}' for i in range(starting_index,starting_index + 65) ]
    elif table == 'Tablets':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'Tablets,{i}' for i in range(starting_index,starting_index + 65) ]
    elif table == 'Mouse':
        if single_run == True:
            return [f'{table},{starting_index}']
        else:
            return [ f'Mouse,{i}' for i in range(starting_index,starting_index + 65) ]

# Exemple d'utilisation
def choose(table_input,filename, starting_id = 0, single_run = None):
    data_list = []
    while data_list == []:
        match table_input:
            case '1':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('Computers', 10000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('Computers', 10000+starting_id)
                    starting_id += 65
            case '2':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('Screens', 20000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('Screens', 20000+starting_id)
                    starting_id += 65
            case '3':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('Phones', 30000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('Phones', 30000+starting_id)
                    starting_id += 65
            case '4':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('Printers', 40000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('Printers', 40000+starting_id)
                    starting_id += 65
            case '5':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('ExternalDrives', 50000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('ExternalDrives', 50000+starting_id)
                    starting_id += 65
            case '6':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('Tablets', 60000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('Tablets', 60000+starting_id)
                    starting_id += 65
            case '7':
                if single_run == True:
                    data_list = gen_qrcodes_bulk('Mouse', 70000+starting_id, single_run = True)
                else:
                    data_list = gen_qrcodes_bulk('Mouse', 70000+starting_id)
                    starting_id += 65
            case _:
                print("Invalid option")
    if single_run == True:
        create_labels_with_qr_codes(data_list, rows, cols, f'wordfiles\\labels_with_pichon{filename}.docx', label_width, label_height, cube_image_path, single_run=True)

    else:
        create_labels_with_qr_codes(data_list, rows, cols, f'wordfiles\\labels_with_pichon{filename}.docx', label_width, label_height, cube_image_path)
    os.system(f'start wordfiles\\labels_with_pichon{filename}.docx')

from urllib.request import urlretrieve

url = 'http://www.samson-agro.com/android-icon-192x192.png'
urlretrieve(url, 'favicon-pichon.png')

rows = 13  # Nombre de rangées d'étiquettes
cols = 5  # Nombre de colonnes d'étiquettes
label_width = 4.05  # Largeur des étiquettes en centimètres
label_height = 2.12  # Hauteur des étiquettes en centimètres
cube_image_path = 'favicon-pichon.png'  # Chemin de l'image du cube

word_files = 'wordfiles/'
if not os.path.exists(word_files):
    os.makedirs(word_files)

qr_path = 'qrcodes/'
if not os.path.exists(qr_path):
    os.makedirs(qr_path)

table_input = 'notdefined'
while table_input not in ['1','2','3','4','5','6','7']:
    print("Choose the type of labels to print :\n1 - Computers\n2 - Screens\n3 - Phones\n4 - Printers\n5 - External Drives\n6 - Tablets\n7 - Mouse")
    table_input = input('Please choose an option : ')

single_run = 'notdefined'
while single_run.lower() != 'y' and single_run.lower() != 'n':
    single_run = input('Print a single label ? :\nY - Yes\nN - No\nPlease choose an option : ')
if single_run.lower() == 'y':
    starting_id = 0
    while starting_id < 1 or starting_id > 9934:
        try:
            starting_id = int(input('Label number (Ex: 52) : '))
        except ValueError:
            pass
    nb_pages = 1
elif single_run.lower() == 'n':
    starting_id = 0
    while starting_id < 1 or starting_id > 9934:
        try:
            starting_id = int(input('From which ID the labels should start ? : '))
        except ValueError:
            pass
            
    nb_pages = 0
    while nb_pages < 1:
        try:
            nb_pages = int(input('Number of pages to print (65 labels per page) : '))
        except ValueError:
            pass

for i in range(1,nb_pages+1):
    if single_run.lower() == 'y':
        choose(table_input, i, starting_id, single_run = True)
    choose(table_input, i, starting_id)
    if i >=1:
        starting_id += 65

# create_labels_with_qr_codes(data_list, rows, cols, 'labels_with_pichon.docx', label_width, label_height, cube_image_path)

# Exemple d'utilisation :

# os.startfile('labels_with_pichon.docx', "print") # Lancement de l'impression sur l'imprimante par défaut

#Cleanup

import glob

os.remove('favicon-pichon.png')

files = glob.glob('qrcodes/*')
for f in files:
    os.remove(f)

os.rmdir('qrcodes/')